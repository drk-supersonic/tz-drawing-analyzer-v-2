"""
core.py — вся логика анализа (без Streamlit).
Импортируется из app.py и run.py.
"""


import json
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

import fitz  # pymupdf
import requests
from docx import Document

# ════════════════════════════════════════════════════════════════
# КОНСТАНТЫ
# ════════════════════════════════════════════════════════════════

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_REFERER = "https://github.com/drk-supersonic/tz-drawing-analyzer"
MODEL = "google/gemini-2.5-flash"

INPUTS_DIR = Path("inputs")
DRAWINGS_DIR = INPUTS_DIR / "chertezhi"

# Список параметров для проверки — используется в обоих запросах
PARAMETERS = [
    "Размеры здания в осях (м)",
    "Высота от чистого пола до верха парапета (м)",
    "Антресольный этаж — отметка уровня",
    "Абсолютная отметка ±0.000 (Балтийская система, м)",
    "Уровень ответственности здания",
    "Степень огнестойкости",
    "Класс конструктивной пожарной опасности (КПО)",
    "Тип каркаса",
    "Сечение колонн (мм × мм)",
    "Шаг колонн (м × м)",
    "Кровля: тип (плоская/скатная)",
    "Кровля: уклон (%)",
    "Кровля: гидроизоляционная мембрана (марка)",
    "Кровля: утеплитель (марка и суммарная толщина мм)",
    "Кровля: водоотвод (внутренний/наружный)",
    "Наружные стеновые панели: толщина (мм)",
    "Наружные стеновые панели: основной цвет RAL",
    "Ворота: тип",
    "Ворота: материал",
    "Ворота: размеры (мм × мм)",
    "Ворота: цвет RAL",
    "Окна: материал профиля",
    "Окна: тип стеклопакета",
    "Окна: цвет RAL",
    "Климат: снеговой район и нагрузка (кг/м²)",
    "Климат: ветровое давление (кг/м²)",
    "Климат: температура наиболее холодной пятидневки (°С)",
]


# ════════════════════════════════════════════════════════════════
# 1. ИЗВЛЕЧЕНИЕ ТЕКСТА
# ════════════════════════════════════════════════════════════════

def docx_to_text(docx_path: str) -> str:
    """Читает DOCX и возвращает чистый текст."""
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def pdf_to_text(pdf_path: str) -> str:
    """
    Извлекает текст из PDF постранично, с разбивкой по блокам.
    Формат: [Страница N] / [блок M] текст — даёт модели точный контекст,
    откуда взято конкретное число, что снижает "фантомные" значения.
    """
    doc = fitz.open(pdf_path)
    pages = []
    for i, page in enumerate(doc):
        # get_text("blocks") -> (x0,y0,x1,y1,text,block_no,block_type)
        # Сортируем по вертикали для правильного порядка чтения сверху вниз
        blocks = sorted(page.get_text("blocks"), key=lambda b: (round(b[1] / 10), b[0]))
        page_lines = []
        for j, block in enumerate(blocks):
            if block[6] != 0:          # block_type 0 = текст; пропускаем картинки
                continue
            text = " ".join(block[4].split()).strip()
            if text:
                page_lines.append(f"  [{j+1}] {text}")
        if page_lines:
            pages.append(f"[Страница {i+1}]\n" + "\n".join(page_lines))
    doc.close()
    return "\n\n".join(pages)


# ════════════════════════════════════════════════════════════════
# 2. ВСПОМОГАТЕЛЬНЫЙ ВЫЗОВ LLM
# ════════════════════════════════════════════════════════════════

def call_llm(api_key: str, system_prompt: str, user_prompt: str, max_tokens: int = 4000,
             _retry: int = 0) -> str:
    """
    Единая точка вызова LLM через OpenRouter.
    max_tokens: extraction=3000, comparison=8000.
    При сетевой ошибке или невалидном ответе — до 3 повторных попыток.
    """
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": OPENROUTER_REFERER,
        "X-Title": "TZ Drawing Analyzer v2",
    }
    payload = {
        "model": MODEL,
        "max_tokens": max_tokens,
        "temperature": 0,  # детерминированный режим — убирает случайные галлюцинации
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    try:
        resp = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=300)
        resp.raise_for_status()
        text = resp.json()["choices"][0]["message"]["content"].strip()
    except (requests.exceptions.RequestException, KeyError) as e:
        if _retry < 3:
            import time
            time.sleep(2 ** _retry)  # экспоненциальная задержка: 1, 2, 4 сек
            return call_llm(api_key, system_prompt, user_prompt, max_tokens, _retry + 1)
        raise RuntimeError(f"API недоступен после 3 попыток: {e}") from e

    # Убираем markdown-обёртки если модель добавила
    if text.startswith("```"):
        text = text.split("\n", 1)[1] if "\n" in text else text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()


def repair_json(text: str) -> str:
    """
    Пытается починить обрезанный или слегка сломанный JSON.
    Стратегии (в порядке применения):
      1. Парсим как есть.
      2. Находим первый '{' и последний '}' — берём только это.
      3. Если обрезан на середине — дорезаем незакрытые скобки/кавычки.
    """
    text = text.strip()

    # Убираем markdown-обёртки
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:] if lines[0].startswith("```") else lines)
    if text.endswith("```"):
        text = text[:-3].rstrip()

    # Попытка 1 — как есть
    try:
        json.loads(text)
        return text
    except json.JSONDecodeError:
        pass

    # Попытка 2 — вырезаем от первого { до последнего }
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        candidate = text[start:end + 1]
        try:
            json.loads(candidate)
            return candidate
        except json.JSONDecodeError:
            pass

    # Попытка 3 — закрываем незакрытые структуры
    # Убираем последний неполный элемент (обрезанная строка)
    candidate = text[start:] if start != -1 else text
    # Обрезаем по последней запятой перед незакрытым объектом
    # Ищем последнюю "чистую" позицию — после закрытого объекта/массива
    for cut_char in [",\n", ", \n", ","]:
        last_comma = candidate.rfind(cut_char)
        if last_comma != -1:
            trimmed = candidate[:last_comma]
            # Считаем незакрытые скобки
            opens = trimmed.count("{") - trimmed.count("}")
            arr_opens = trimmed.count("[") - trimmed.count("]")
            # Добавляем закрывающие скобки
            closing = ("]" * max(0, arr_opens)) + ("}" * max(0, opens))
            fixed = trimmed + closing
            try:
                json.loads(fixed)
                return fixed
            except json.JSONDecodeError:
                continue

    # Не починили — вернём как есть, parse_json_response выдаст понятную ошибку
    return text


def parse_json_response(text: str, context: str = "",
                        api_key: str = "", system: str = "", user: str = "",
                        max_tokens: int = 4000) -> dict:
    """
    Парсит JSON из ответа LLM.
    Если не удалось — пробует починить через repair_json.
    Если всё равно не удалось и переданы параметры запроса — повторяет вызов API
    с явной инструкцией вернуть валидный JSON.
    """
    repaired = repair_json(text)
    try:
        return json.loads(repaired)
    except json.JSONDecodeError:
        pass

    # Если есть параметры запроса — повторяем с усиленной инструкцией
    if api_key and user:
        retry_system = system + "\n\nКРИТИЧНО: верни ТОЛЬКО валидный JSON. Никакого текста до или после. Никаких пояснений."
        retry_user = user + "\n\nПОВТОРНЫЙ ЗАПРОС: предыдущий ответ содержал невалидный JSON. Верни только JSON, ничего более."
        try:
            retry_text = call_llm(api_key, retry_system, retry_user, max_tokens)
            retry_repaired = repair_json(retry_text)
            return json.loads(retry_repaired)
        except (json.JSONDecodeError, RuntimeError):
            pass

    raise RuntimeError(
        f"JSON parse error ({context}) — не удалось исправить даже после повтора.\n"
        f"Ответ модели (первые 1500 символов):\n{text[:1500]}"
    )


# ════════════════════════════════════════════════════════════════
# 3. ШАГ 1 — ИЗВЛЕЧЕНИЕ ПАРАМЕТРОВ ИЗ ДОКУМЕНТА
# ════════════════════════════════════════════════════════════════

def extract_parameters(api_key: str, source_name: str, text: str, source_type: str) -> dict:
    """
    Запрос 1: Извлекает структурированные параметры из одного документа.

    source_type: "tz" или "drawing"
    Возвращает: {"parameters": {"Название параметра": {"value": "...", "source": "...", "raw": "..."}}}
    """
    params_list = "\n".join(f"- {p}" for p in PARAMETERS)

    if source_type == "tz":
        doc_description = "Техническое Задание (ТЗ) на строительство"
        instruction = (
            "Извлеки значения всех параметров, которые явно указаны в ТЗ. "
            "Если параметр в ТЗ не указан — укажи null. "
            "Не интерпретируй и не вычисляй — только то что написано дословно."
        )
    else:
        doc_description = f"Рабочий чертёж: {source_name}"
        instruction = """Правила извлечения из чертежа:
- Текст разбит по блокам: [Страница N] / [блок M]. Используй это для точного указания источника.
- Извлекай значение ТОЛЬКО если оно явно относится к данному параметру по контексту блока.
- СТОП-список — не путай похожие числа из разных контекстов:
  * "Шаг колонн" — только явный шаг несущих колонн основного каркаса; не шаг ригелей, не шаг панелей, не разбивочная сетка
  * "Сечение колонн" — только сечение несущих колонн каркаса, не второстепенных элементов
  * "Толщина наружных стеновых панелей" — только наружные ограждающие панели, не перегородки и не внутренние стены
  * "Высота до верха парапета" — от уровня чистого пола (±0.000) до верхней отметки парапета
  * "Абсолютная отметка ±0.000" — геодезическая отметка уровня чистого пола в Балтийской системе
- Для многослойных конструкций (утеплитель из нескольких слоёв): все слои в "raw", сумму в "value"
- Если параметр явно не указан или контекст неоднозначен — верни null, не угадывай
- В поле "source" указывай: номер страницы и содержательное название раздела/таблицы"""

    system = (
        "Ты эксперт по строительной документации. "
        "Отвечай ТОЛЬКО валидным JSON без пояснений и markdown."
    )

    user = f"""Документ: {doc_description}

{instruction}

Параметры для извлечения:
{params_list}

Верни JSON в формате:
{{
  "parameters": {{
    "Название параметра": {{
      "value": "извлечённое значение или null",
      "source": "страница/раздел где найдено",
      "raw": "дословная цитата из документа"
    }}
  }}
}}

ТЕКСТ ДОКУМЕНТА:
{text}"""

    response_text = call_llm(api_key, system, user, max_tokens=3000)
    return parse_json_response(response_text, context=f"extraction:{source_name}", api_key=api_key, system=system, user=user, max_tokens=3000)



# ════════════════════════════════════════════════════════════════
# 3b. ШАГ 1c — ВЕРИФИКАЦИЯ ИЗВЛЕЧЁННЫХ ЗНАЧЕНИЙ
# ════════════════════════════════════════════════════════════════

def verify_extracted(api_key: str, drawings_data: dict, tz_data: dict) -> dict:
    """
    Шаг 1c: Проверяет извлечённые из чертежей значения на контекстную корректность.
    Убирает "фантомные" значения — числа, попавшие из неподходящего контекста.

    Пример: шаг колонн "11x18; 6x6" → модель оставляет только "11x18",
    потому что 6x6 — это шаг другого элемента, не несущих колонн каркаса.
    """
    drawings_json = json.dumps(drawings_data, ensure_ascii=False, indent=2)
    tz_json = json.dumps(tz_data, ensure_ascii=False, indent=2)

    system = (
        "Ты эксперт по строительной документации. "
        "Отвечай ТОЛЬКО валидным JSON без пояснений и markdown."
    )

    user = f"""Проверь извлечённые из чертежей значения параметров на корректность.

Тебе переданы:
1. Данные из ТЗ (для контекста — какие значения ожидаются)
2. Данные, извлечённые из чертежей

## ДАННЫЕ ТЗ (только для контекста):
{tz_json}

## ДАННЫЕ ИЗ ЧЕРТЕЖЕЙ (требуют проверки):
{drawings_json}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ЗАДАЧА
━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Для каждого параметра каждого чертежа проверь: не попало ли в value значение
из постороннего контекста? Если в value несколько значений через ";", определи
какое из них действительно относится к данному параметру.

КРИТЕРИИ "фантомного" значения (удали его):
- "Шаг колонн": любое значение кроме явного шага НЕСУЩИХ КОЛОНН ОСНОВНОГО КАРКАСА.
  Шаг 6x6, 3x3 и т.п. — это шаг ригелей, плит или разбивочная сетка, не колонн.
- "Сечение колонн": только сечение основных несущих колонн, не стоек и не пилонов.
- "Толщина наружных стеновых панелей": только наружные панели, не перегородки.
- Любой параметр с цветом RAL: если несколько значений RAL — оставь то, которое
  явно относится к данному типу конструкции (панели, ворота, окна), остальные удали.
  Исключение: если разные значения из разных чертежей (это реальное противоречие — оставь оба).
- Числа из таблиц спецификаций (количество штук, площадь, масса) — не являются
  значениями конструктивных параметров.

ПРАВИЛО: если после удаления фантомов остаётся одно значение — верни его как строку.
Если остаётся несколько реально разных значений из разных файлов — верни через "; ".
Если значение корректно — оставь без изменений.
Если нет уверенности — ОБЯЗАТЕЛЬНО оставь исходное значение без изменений.
ПРАВИЛО КОНСЕРВАТИЗМА: лучше оставить лишнее значение, чем удалить нужное.
Удаляй только когда АБСОЛЮТНО очевидно что значение из чужого контекста.

Верни JSON той же структуры что получил в "ДАННЫЕ ИЗ ЧЕРТЕЖЕЙ", только с исправленными value:
{{
  "имя_файла.pdf": {{
    "parameters": {{
      "Название параметра": {{
        "value": "проверенное значение или null",
        "source": "без изменений",
        "raw": "без изменений"
      }}
    }}
  }}
}}"""

    response_text = call_llm(api_key, system, user, max_tokens=4000)
    verified = parse_json_response(response_text, context="verify_extracted", api_key=api_key, system=system, user=user, max_tokens=4000)

    # Если верификация вернула неполный результат — используем оригинал
    for file_name in drawings_data:
        if file_name not in verified:
            verified[file_name] = drawings_data[file_name]
        else:
            # Для каждого параметра: если верификация вернула null но оригинал не null — берём оригинал
            orig_params = drawings_data[file_name].get("parameters", {})
            verif_params = verified[file_name].get("parameters", {})
            for param, orig_val in orig_params.items():
                if param not in verif_params:
                    verif_params[param] = orig_val

    return verified

# ════════════════════════════════════════════════════════════════
# 4. ШАГ 2 — СРАВНЕНИЕ СТРУКТУРИРОВАННЫХ ДАННЫХ
# ════════════════════════════════════════════════════════════════

def compare_parameters(
    api_key: str,
    tz_data: dict,
    drawings_data: Dict[str, dict],
    drawing_names: list
) -> dict:
    """
    Запрос 2: Сравнивает уже извлечённые структурированные данные.
    На вход — чистые JSON-объекты, не сырой текст.
    Правила нормализации универсальные — без привязки к конкретным документам.
    """

    tz_json = json.dumps(tz_data, ensure_ascii=False, indent=2)
    drawings_json = json.dumps(drawings_data, ensure_ascii=False, indent=2)

    system = (
        "Ты эксперт по строительному контролю. "
        "Отвечай только на русском языке. "
        "Возвращай ТОЛЬКО валидный JSON, без пояснений и markdown."
    )

    user = f"""Сравни параметры Технического Задания с параметрами рабочих чертежей.

Тебе переданы уже извлечённые структурированные данные. Твоя задача — только сравнить значения.

## ДАННЫЕ ТЕХНИЧЕСКОГО ЗАДАНИЯ:
{tz_json}

## ДАННЫЕ ЧЕРТЕЖЕЙ (по файлам):
{drawings_json}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ПРАВИЛА НОРМАЛИЗАЦИИ (применяй перед сравнением)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Перед сравнением нормализуй оба значения по следующим правилам:

1. ЧИСЛА И ЕДИНИЦЫ
   - Запятая и точка как разделитель дробной части эквивалентны: 3,600 = 3.600
   - Пробел между числом и единицей несущественен: "30 мм" = "30мм" = "30 mm"
   - Знак минус и слово "минус" эквивалентны: "-25°С" = "минус 25°С"
   - Сумма слоёв/компонентов: если в чертеже несколько слоёв одного материала,
     сравнивай их сумму с суммарным значением в ТЗ. Пример: "40мм + 100мм" → 140мм

2. ФОРМАТЫ ЗАПИСИ РАЗМЕРОВ
   - Разделители "×", "x", "х" (латинская/кириллическая), "*" эквивалентны
   - Порядок следования WxH = HxW для симметричных зданий
   - "108,0 × 33,0" = "108x33" = "108.0 х 33.0"

3. УТОЧНЕНИЯ БЕЗ ПРОТИВОРЕЧИЯ
   - Если чертёж содержит уточнение, которое не противоречит ТЗ — это совпадение.
     Примеры: "металлические утеплённые" при ТЗ "металлические",
              "II (нормальный)" при ТЗ "II",
              "подъёмно-секционные автоматические" при ТЗ "подъёмно-секционные"

4. КИРИЛЛИЦА / ЛАТИНИЦА
   - Одинаковые аббревиатуры и марки в разных алфавитах эквивалентны:
     "RAL" = "РАЛ", "ПВХ" = "PVC"

5. КОСВЕННОЕ ПОДТВЕРЖДЕНИЕ
   - Если параметр явно присутствует в чертеже (план на отметке +X.XXX
     подтверждает уровень +X.XXX) — это совпадение

6. ЧАСТИЧНОЕ СОВПАДЕНИЕ ПРИ МНОЖЕСТВЕННЫХ ЗНАЧЕНИЯХ
   - Если в чертеже извлечено несколько значений через ";" и ХОТЯ БЫ ОДНО совпадает
     с ТЗ после нормализации — это СОВПАДЕНИЕ, не расхождение.
   - Логика: второе значение могло попасть из другого контекста документа (другой элемент,
     другой узел), но основной параметр подтверждён.
   - Пример: ТЗ "11,0 × 18,0", чертёж "11.0 x 18.0; 6x6" → совпадение (11x18 совпадает)
   - Пример: ТЗ "RAL 7016", чертёж "RAL 7016; RAL 7004" → совпадение (7016 совпадает)

7. ДИАПАЗОН В ТЗ — ОДНО ЗНАЧЕНИЕ В ЧЕРТЕЖЕ
   - Если ТЗ содержит диапазон или два значения через "/" или "—",
     а чертёж показывает одно из этих значений — это СОВПАДЕНИЕ.
   - Пример: ТЗ "-25°С / -29°С", чертёж "-29" → совпадение (входит в диапазон)
   - Пример: ТЗ "III — 180 кг/м²", чертёж "III - 180" → совпадение

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ПРАВИЛА КЛАССИФИКАЦИИ
━━━━━━━━━━━━━━━━━━━━━━━━━━━━

После нормализации отнеси каждый параметр к одной категории:

- "matches" — значения эквивалентны после нормализации (включая правила 6 и 7 выше)
- "discrepancies" — значения реально отличаются по смыслу после применения ВСЕХ правил.
  НЕЛЬЗЯ создавать расхождение если в ТЗ значение null.
  НЕЛЬЗЯ создавать расхождение если хотя бы одно из значений чертежа совпадает с ТЗ (правило 6).
  НЕЛЬЗЯ создавать расхождение если значение чертежа входит в диапазон ТЗ (правило 7).
- "missing_in_drawings" — значение в ТЗ НЕ null И не найдено ни в одном чертеже.
  СТРОГО ЗАПРЕЩЕНО: помещать сюда параметр если в ТЗ значение null — только в extra_in_drawings.
- "extra_in_drawings" — значение в ТЗ null (не указано), но найдено в чертежах.
  Если параметр здесь — он НЕ должен быть в missing_in_drawings ни при каких условиях.

ПРАВИЛА ОБЪЕДИНЕНИЯ ИСТОЧНИКОВ:
- Если параметр найден в нескольких чертежах с одинаковым значением (после нормализации):
  → одна строка в нужной категории, источники через точку с запятой. НЕ дублируй строки.
- Если значения в разных чертежах реально противоречат друг другу И НИ ОДНО не совпадает с ТЗ:
  → одна строка в "discrepancies", в поле "drawing" укажи оба значения через " / ",
    в "source" — оба источника. НЕ создавай отдельную строку на каждый чертёж.
- Итого: каждый параметр присутствует в итоговом JSON ровно в одной строке одной категории.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ФОРМАТ ОТВЕТА
━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{{
  "matches": [
    {{"parameter": "название", "requirement": "значение по ТЗ", "drawing": "значение в чертежах", "source": "файл, страница"}}
  ],
  "discrepancies": [
    {{"parameter": "название", "requirement": "значение по ТЗ", "drawing": "значение в чертежах", "criticality": "КРИТИЧНО/ВАЖНО/НЕЗНАЧИТЕЛЬНО", "source": "файл, страница", "comment": "пояснение"}}
  ],
  "missing_in_drawings": [
    {{"parameter": "название", "requirement": "значение по ТЗ", "criticality": "КРИТИЧНО/ВАЖНО/НЕЗНАЧИТЕЛЬНО"}}
  ],
  "extra_in_drawings": [
    {{"parameter": "название", "drawing": "значение в чертежах", "source": "файл, страница", "comment": "пояснение"}}
  ],
  "summary": {{
    "total_matches": 0,
    "total_discrepancies": 0,
    "total_missing": 0,
    "total_extra": 0,
    "recommendations": ["рекомендация"]
  }}
}}

Уровни критичности расхождений:
- КРИТИЧНО: влияет на объём/привязку здания, безопасность, нормативные требования
  (высота здания, абсолютные отметки, несущие конструкции, противопожарные классы)
- ВАЖНО: влияет на функциональность, теплотехнику, эксплуатацию
  (толщина ограждающих конструкций, тип утеплителя, инженерные системы)
- НЕЗНАЧИТЕЛЬНО: уточняющие характеристики, косметические различия, формат записи"""

    response_text = call_llm(api_key, system, user, max_tokens=8000)
    result = parse_json_response(response_text, context="comparison", api_key=api_key, system=system, user=user, max_tokens=8000)

    # Пересчитываем счётчики (не доверяем модели)
    result["summary"]["total_matches"] = len(result.get("matches", []))
    result["summary"]["total_discrepancies"] = len(result.get("discrepancies", []))
    result["summary"]["total_missing"] = len(result.get("missing_in_drawings", []))
    result["summary"]["total_extra"] = len(result.get("extra_in_drawings", []))

    return result


# ════════════════════════════════════════════════════════════════
# 5. ОРКЕСТРАТОР — ДВУХЭТАПНЫЙ ПАЙПЛАЙН
# ════════════════════════════════════════════════════════════════

def run_analysis(
    api_key: str,
    tz_text: str,
    drawings: Dict[str, str],
    progress_callback=None,
) -> Dict[str, Any]:
    """
    Основной пайплайн:
      1. Извлекаем параметры из ТЗ
      2. Извлекаем параметры из каждого чертежа
      3. Сравниваем структурированные данные
    """

    def update_progress(pct: int, msg: str):
        if progress_callback:
            progress_callback(pct, msg)

    # ── Шаг 1a: извлечение из ТЗ ──────────────────────────────────
    update_progress(15, "Извлекаю параметры из Технического Задания...")
    tz_data = extract_parameters(api_key, "ТЗ", tz_text, source_type="tz")

    # ── Шаг 1b: извлечение из каждого чертежа ─────────────────────
    drawings_data = {}
    n = len(drawings)
    for idx, (name, text) in enumerate(drawings.items()):
        pct = 30 + int(35 * (idx / n))
        update_progress(pct, f"Извлекаю параметры из чертежа {idx+1}/{n}: {name}...")
        drawings_data[name] = extract_parameters(api_key, name, text, source_type="drawing")

    # ── Шаг 1c: верификация — убираем фантомные значения ───────────
    update_progress(68, "Верифицирую извлечённые значения...")
    drawings_data = verify_extracted(api_key, drawings_data, tz_data)

    # ── Шаг 2: сравнение структурированных данных ──────────────────
    update_progress(78, "Сравниваю параметры ТЗ и чертежей...")
    result = compare_parameters(
        api_key,
        tz_data,
        drawings_data,
        drawing_names=list(drawings.keys()),
    )

    # Сохраняем промежуточные данные для отладки
    result["_debug"] = {
        "tz_extracted": tz_data,
        "drawings_extracted": drawings_data,
    }

    return result


# ════════════════════════════════════════════════════════════════
# 6. ГЕНЕРАТОР ОТЧЁТА
# ════════════════════════════════════════════════════════════════

def generate_report(result: Dict[str, Any], tz_name: str, pdf_names: list) -> str:
    now = datetime.now().strftime("%d.%m.%Y %H:%M")
    s = result["summary"]
    lines = []
    lines.append("# Отчёт о соответствии чертежей техническому заданию\n")
    lines.append("## Общая информация")
    lines.append(f"- Дата анализа: {now}")
    lines.append(f"- Техническое задание: `{tz_name}`")
    lines.append(f"- Проанализировано чертежей: {len(pdf_names)}")
    for n in pdf_names:
        lines.append(f"  - `{n}`")
    lines.append(f"- Совпадений: **{s['total_matches']}**")
    lines.append(f"- Расхождений: **{s['total_discrepancies']}**")
    lines.append(f"- Отсутствует в чертежах: **{s['total_missing']}**")
    lines.append("")

    if result.get("matches"):
        lines.append("## Совпадения\n")
        lines.append("| Параметр | ТЗ | Чертежи | Источник |")
        lines.append("|----------|-----|---------|----------|")
        for m in result["matches"]:
            lines.append(
                f"| {m.get('parameter','')} "
                f"| {m.get('requirement','')} "
                f"| {m.get('drawing','')} "
                f"| {m.get('source','')} |"
            )
        lines.append("")

    if result.get("discrepancies"):
        lines.append("## Расхождения\n")
        lines.append("| Параметр | ТЗ | Чертежи | Критичность | Источник | Комментарий |")
        lines.append("|----------|-----|---------|-------------|----------|-------------|")
        for d in result["discrepancies"]:
            lines.append(
                f"| {d.get('parameter','')} "
                f"| {d.get('requirement','')} "
                f"| {d.get('drawing','')} "
                f"| {d.get('criticality','')} "
                f"| {d.get('source','')} "
                f"| {d.get('comment','')} |"
            )
        lines.append("")

    if result.get("missing_in_drawings"):
        lines.append("## Отсутствует в чертежах\n")
        lines.append("| Параметр | ТЗ | Критичность |")
        lines.append("|----------|----|-------------|")
        for m in result["missing_in_drawings"]:
            lines.append(
                f"| {m.get('parameter','')} "
                f"| {m.get('requirement','')} "
                f"| {m.get('criticality','')} |"
            )
        lines.append("")

    if result.get("extra_in_drawings"):
        lines.append("## Дополнительные данные из чертежей\n")
        lines.append("> Эти параметры присутствуют в чертежах, но не указаны в ТЗ. Расхождением не являются.")
        lines.append("")
        lines.append("| Параметр | Значение в чертежах | Источник | Комментарий |")
        lines.append("|----------|---------------------|----------|-------------|")
        for e in result["extra_in_drawings"]:
            lines.append(
                f"| {e.get('parameter','')} "
                f"| {e.get('drawing','')} "
                f"| {e.get('source','')} "
                f"| {e.get('comment','')} |"
            )
        lines.append("")

    lines.append("## Выводы\n")
    lines.append(f"- Итого совпадений: {s['total_matches']}")
    lines.append(f"- Итого расхождений: {s['total_discrepancies']}")
    lines.append(f"- Отсутствует в чертежах: {s['total_missing']}")
    lines.append("")
    if s.get("recommendations"):
        lines.append("**Рекомендации:**")
        for r in s["recommendations"]:
            lines.append(f"- {r}")

    return "\n".join(lines)


# ════════════════════════════════════════════════════════════════
